# -*- coding: utf-8 -*-
"""详细版问题一二三论文：公式衔接说明 + 可视化插图预留。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_DIR = Path(__file__).resolve().parent
OUT = DOC_DIR / "2025B论文_问题一二三_详细版.docx"
MD_OUT = DOC_DIR / "2025B论文_问题一二三_详细版.md"
BACKUP = DOC_DIR / "2025B论文_填充前备份.docx"


def set_run_font(run, *, east_asia="宋体", size_pt=10.5, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def style_p(p, *, level="body", first_indent=True):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    if level == "title":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Pt(0)
    elif level in {"h1", "h2", "h3", "caption", "formula", "figure_box"}:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if level in {"caption", "formula", "figure_box", "title"} else WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Pt(0)
        if level == "h1":
            pf.space_before = Pt(12)
        elif level == "h2":
            pf.space_before = Pt(8)
        if level != "formula" and level != "caption" and level != "figure_box":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.74) if first_indent else Pt(0)


def add_para(doc, text, *, level="body", first_indent=True):
    p = doc.add_paragraph()
    style_p(p, level=level, first_indent=first_indent)
    if level == "title":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=15, bold=True)
    elif level == "h1":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=15, bold=True)
    elif level == "h2":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=12, bold=True)
    elif level == "h3":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=10.5, bold=True)
    elif level == "caption":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=10.5, bold=True)
    elif level == "figure_box":
        set_run_font(p.add_run(text), east_asia="宋体", size_pt=10.5, color=RGBColor(0x80, 0x00, 0x00))
    else:
        set_run_font(p.add_run(text), east_asia="宋体", size_pt=10.5)
    return p


def add_formula(doc, text, tag: str):
    add_para(doc, f"{text}　　({tag})", level="formula")


def add_follow(doc, text):
    """公式/图后接排解释，首行不缩进。"""
    add_para(doc, text, first_indent=False)


def add_figure_slot(doc, fig_no: str, title: str, rel_path: str, note: str = ""):
    """预留插图位置，便于手工插入项目生成的 PNG。"""
    box = (
        f"【插图预留：图{fig_no}】\n"
        f"请插入图像文件：{rel_path}\n"
        f"建议版式：居中，宽度约 12–14 cm"
    )
    if note:
        box += f"\n读图要点：{note}"
    p = add_para(doc, box, level="figure_box")
    # 加边框提示占位
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "C00000")
        pBdr.append(el)
    pPr.append(pBdr)
    add_para(doc, f"图{fig_no}  {title}", level="caption")


def add_table(doc, headers, rows, caption):
    add_para(doc, caption, level="caption")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), east_asia="黑体", size_pt=10.5, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(val), east_asia="宋体", size_pt=10.5)
    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_para(doc, "碳化硅外延层厚度的红外干涉测量及多光束干涉影响分析", level="title")

    # ========== 一、问题重述 ==========
    add_para(doc, "一、问题重述", level="h1")
    add_para(doc, "1.1 问题背景", level="h2")
    add_para(
        doc,
        "碳化硅（SiC）是典型的第三代宽禁带半导体材料，外延层厚度直接关系到击穿电压、导通电阻等器件关键指标。"
        "红外干涉测厚利用空气–外延层与外延层–衬底两界面反射光的干涉条纹，由条纹周期反演膜厚，具有非接触、无损的优点。"
        "然而实测光谱同时受到强吸收带、慢变基线、仪器标定误差以及可能的多次反射影响；若仍用理想正入射、常折射率、纯余弦条纹模型，"
        "容易引入系统偏差。因此，需要把斜入射几何、条纹提取策略与多光束可观测判定统一进可检验的建模流程。",
    )
    add_figure_slot(
        doc,
        "1",
        "数学模型总体流程图（请插入项目流程图）",
        "output/model_flowchart.png",
        "用于总览：预处理→双光束→多光束诊断→门控输出",
    )

    add_para(doc, "1.2 问题要求", level="h2")
    add_para(
        doc,
        "题目给出四份红外光谱附件：附件1、2为同一碳化硅晶圆在入射角 10°、15° 下的反射谱；"
        "附件3、4为硅片在 10°、15° 下的反射谱。每份约 7469 点，波数约 399.7–4000.1 cm⁻¹，间隔约 0.482 cm⁻¹。"
        "本文按三问递进展开。",
    )
    add_para(
        doc,
        "问题一：在“每个界面仅发生一次反射/透射”的双光束假设下，建立斜入射薄膜干涉厚度模型，"
        "写出折射角、光程差、相位差，并得到厚度与相邻同类极值波数间隔的解析关系。",
    )
    add_para(
        doc,
        "问题二：利用附件1、2计算同一块碳化硅晶圆外延层厚度；需兼顾透明波段选取、去噪去基线、峰谷稳健估计，"
        "并以双角度一致性与不确定度评价可靠性。",
    )
    add_para(
        doc,
        "问题三：给出多光束干涉可观测条件；判断附件3、4是否存在显著多光束效应并测厚；"
        "同时检验附件1、2是否需要多光束修正，避免“残差下降即判定多光束”的误判。",
    )

    # ========== 二、问题分析 ==========
    add_para(doc, "二、问题分析", level="h1")
    add_para(
        doc,
        "三问共享同一光学主线：由几何光学得到相位差 δ，由 δ 的极值间隔得到厚度 d；"
        "差别在于问题一止于解析关系，问题二把关系落到 SiC 数据与可靠性，问题三把模型从双光束扩展到多光束并设置证据门控。",
    )
    add_para(doc, "2.1 问题一分析", level="h2")
    add_para(
        doc,
        "关键是把斜入射正确写入相位。Snell 定律决定膜内折射角 θ₁，从而决定往返光程差 ΔL=2nd cosθ₁；"
        "相位差 δ= (2π/λ)ΔL + φᵣ。相邻同类极值对应相位改变 2π，于是厚度由波数间隔 Δν̃ 唯一确定。"
        "该公式在 θ₀→0 时应退化为正入射公式，可作为推导自洽检验。问题一不直接给数值，而是为后两问提供统一锚点。",
    )
    add_para(doc, "2.2 问题二分析", level="h2")
    add_para(
        doc,
        "难点在于全谱并非处处满足常折射率透明假设：SiC 在约 797–1000 cm⁻¹ 有强声子吸收，会扭曲条纹。"
        "策略是：先自动选择透明窗（本项目最终主波段约 1200–4000 cm⁻¹），再双尺度滤波分离慢变基线与快变条纹；"
        "用 FFT 粗估厚度，用峰/谷 Theil–Sen 稳健回归作为主估计，并用间距 bootstrap 给条件置信区间；"
        "最后以两角度相对差与逆方差加权融合给出报告厚度。",
    )
    add_para(doc, "2.3 问题三分析", level="h2")
    add_para(
        doc,
        "多光束使条纹偏离纯余弦并引入高次谐波。若仅比较 RMSE，多参数 Airy 模型几乎总会“更好”，因此必须设置联合证据："
        "谐波比 A₂/A₁、有效反射率 Rᵢ、RMSE 改善与 ΔAICc。四项同时通过才采用多光束厚度；否则回退双光束。"
        "据此，硅片与碳化硅可能走向不同模型分支，这正是问题三相对问题二的结构扩展。",
    )

    # ========== 三、假设 ==========
    add_para(doc, "三、模型假设", level="h1")
    for i, s in enumerate(
        [
            "外延层上下表面局部平行，分析光斑内厚度可视为常数。",
            "空气折射率 n₀=1，入射角相对表面法线定义。",
            "主厚度反演在弱吸收透明波段进行；SiC 强声子带不参与主估计。",
            "双/多光束条纹阶段取 n_SiC=2.55、n_Si=3.42；结果为该光学常数下的条件估计。",
            "基线与振幅可慢变，但同一材料两入射角共享同一物理厚度。",
            "多光束判定必须四项证据同时通过，禁止仅凭残差下降改判模型。",
        ],
        1,
    ):
        add_para(doc, f"{i}. {s}")

    # ========== 四、符号 ==========
    add_para(doc, "四、符号说明", level="h1")
    add_para(doc, "主要符号见表1。后文公式编号与符号保持一致。")
    add_table(
        doc,
        ["符号", "含义", "单位"],
        [
            ["θ₀, θ₁", "入射角、膜内折射角", "°"],
            ["n", "外延层折射率（常折射率阶段）", "—"],
            ["ν̃, λ", "波数、波长", "cm⁻¹, µm"],
            ["d", "外延层厚度", "µm"],
            ["ΔL", "膜内一次往返光程差", "cm"],
            ["δ, φᵣ", "相位差、界面反射附加相位", "rad"],
            ["g", "光学相位坐标 ν̃ n cosθ₁", "cm⁻¹"],
            ["Δν̃", "相邻同类极值波数间隔", "cm⁻¹"],
            ["A,B", "慢变基线与条纹振幅", "%"],
            ["Rᵢ", "Airy 有效界面反射率", "—"],
            ["F", "Airy 精细度系数", "—"],
            ["f₁, A₂/A₁", "基频、二次谐波比", "cycles/cm⁻¹, —"],
            ["ΔAICc", "双光束相对多光束的 AICc 差", "—"],
        ],
        "表1  主要符号说明",
    )

    # ========== 五、问题一 ==========
    add_para(doc, "五、问题一模型的建立与求解", level="h1")
    add_para(doc, "5.1 几何光学：由折射到光程差", level="h2")
    add_para(
        doc,
        "设红外光自空气以入射角 θ₀ 射入折射率为 n 的外延层，折射角为 θ₁。Snell 定律给出",
    )
    add_formula(doc, "sin θ₀ = n sin θ₁", "1")
    add_follow(
        doc,
        "式(1)把空气侧可测的入射角映射到膜内传播方向。对无吸收实折射率，由勾股关系得到",
    )
    add_formula(doc, "cos θ₁ = √(1 − sin²θ₀ / n²)", "2")
    add_follow(
        doc,
        "式(2)要求 n > sinθ₀，否则发生全反射；本文附件角度与折射率均满足该条件。"
        "光束在上表面反射一次、在下界面反射一次后回到上表面，膜内几何路径长度为 2d/cosθ₁，"
        "但干涉比较的是光程，还需乘折射率，并注意到波前法向投影，等价于常用形式",
    )
    add_formula(doc, "ΔL = 2 n d cos θ₁", "3")
    add_follow(
        doc,
        "式(3)是连接“厚度 d”与“干涉相位”的第一座桥梁：厚度越大，或 n cosθ₁ 越大，往返光程差越大。",
    )

    add_para(doc, "5.2 从光程差到相位差，再到反射率振荡", level="h2")
    add_para(
        doc,
        "相位差等于波数乘以光程差（再计界面反射附加相位）。因 ν̃=1/λ，有",
    )
    add_formula(doc, "δ(ν̃) = (2π/λ) ΔL + φᵣ = 4π d ν̃ n cos θ₁ + φᵣ", "4")
    add_follow(
        doc,
        "式(4)直接由式(3)代入得到：把 ΔL 换成 2nd cosθ₁，并把 2π/λ=2πν̃ 代入即可。"
        "附加相位 φᵣ 主要影响极值对应的干涉级次奇偶，但不改变相邻同类极值之间的间隔，"
        "因此后文用间隔反演厚度时可把它视为常数。定义光学相位坐标",
    )
    add_formula(doc, "g(ν̃,θ₀) = ν̃ n cos θ₁ = ν̃ √(n² − sin²θ₀)", "5")
    add_follow(
        doc,
        "式(5)把式(2)嵌入式(4)：δ=4π d g + φᵣ。于是“在波数轴上走多远对应一个干涉周期”完全由 g 的变化率决定。"
        "在双光束假设下，两束反射场叠加，慢变包络记为 A、B，反射率可写",
    )
    add_formula(doc, "R(ν̃) = A(ν̃) + B(ν̃) cos δ(ν̃)", "6")
    add_follow(
        doc,
        "式(6)说明：快速振荡完全由 δ 控制，而 A、B 吸收仪器基线与振幅慢变。"
        "这正是数值上先去基线、再对残差做 FFT/峰谷分析的理论依据。",
    )

    add_para(doc, "5.3 厚度—波数间隔解析式及其退化检验", level="h2")
    add_para(
        doc,
        "设两个同类极值（同为峰或同为谷）ν̃_a、ν̃_b 之间跨越 Δm 个干涉级次，则相位差改变 2πΔm，即",
    )
    add_formula(doc, "4π d [g(ν̃_b) − g(ν̃_a)] = 2π Δm", "7")
    add_follow(doc, "整理式(7)得到一般色散情形下的厚度公式")
    add_formula(doc, "d = Δm / {2 [g(ν̃_b,θ₀) − g(ν̃_a,θ₀)]}", "8")
    add_follow(
        doc,
        "若窄波段内 n 近似常数，且取相邻同类极值（Δm=1），则 g(ν̃_b)−g(ν̃_a)=Δν̃·n cosθ₁，"
        "式(8)化为问题一核心结果",
    )
    add_formula(doc, "d = 1 / [2 Δν̃ √(n² − sin²θ₀)]", "9")
    add_follow(
        doc,
        "当 d 以 µm、Δν̃ 以 cm⁻¹ 计时，式(9)右侧再乘 10⁴，与程序 optics.thickness_from_fringe_spacing 一致。"
        "令 θ₀→0，式(9)退化为 d=1/(2nΔν̃)，与正入射教材公式一致，说明斜入射推广正确。"
        "至此，问题一完成了“θ₀→θ₁→ΔL→δ→g→d(Δν̃)”的完整公式链。",
    )

    add_para(doc, "5.4 与后续数值实现的接口", level="h2")
    add_para(
        doc,
        "问题一给出的是解析映射。落地时需要从离散光谱估计 Δν̃：先去噪去基线得到残差（对应式(6)中去掉 A 后的振荡项），"
        "再由 FFT 粗估周期，或由峰谷序列的 Theil–Sen 斜率得到稳健 Δν̃，最后代入式(9)。"
        "该接口在问题二中完整实现。",
    )

    # ========== 六、问题二 ==========
    add_para(doc, "六、问题二模型的建立与求解", level="h1")
    add_para(doc, "6.1 数据预处理：把式(6)中的慢变与快变分开", level="h2")
    add_para(
        doc,
        "附件1、2对应同一 SiC 晶圆的 10°、15° 光谱。为使式(9)成立，必须避开强吸收带。"
        "本项目对候选透明窗评分后，主拟合波段取约 1200–4000 cm⁻¹。"
        "短窗 Savitzky–Golay 抑制点噪声，长窗提取慢变基线 A(ν̃)，残差近似 B cosδ，供后续反演。"
        "附件2最大反射率约 102.74%，提示绝对标定可能异常；由于双光束厚度主要由峰位/间距决定，"
        "程序保留原始量程并记入审计，不强制裁剪到 100%。",
    )
    add_figure_slot(
        doc,
        "2",
        "附件1（SiC，10°）光谱处理—拟合完整证据图",
        "output/sic_10_fit.png",
        "关注：拟合波段、去基线条纹、峰谷标注与残差",
    )
    add_figure_slot(
        doc,
        "3",
        "附件2（SiC，15°）光谱处理—拟合完整证据图",
        "output/sic_15_fit.png",
        "与图2对照，检查双角度条纹形态是否一致",
    )

    add_para(doc, "6.2 求解策略：粗估—稳健极值—精修—不确定度", level="h2")
    add_para(
        doc,
        "步骤与式(9)的对应关系如下。"
        "（1）FFT 粗估：残差在波数域的主频 f₁≈2 n cosθ₁ · d · 10⁻⁴，由此得到 d 的初值；"
        "这是式(4)(5)在频域的直接体现。"
        "（2）峰谷提取：按粗周期设最小峰距，分别得到峰序列与谷序列；"
        "对序号–波数做 Theil–Sen 回归，斜率即稳健 Δν̃，代入式(9)得 d_peak、d_valley。"
        "（3）精修：在稳健解邻域对相位模型做有界优化；若色散导致全谱目标跳级，则保留峰谷解。"
        "（4）统计区间：对峰/谷间距重采样，得到条件 95% 置信区间；"
        "系统误差则由 d∝(n²−sin²θ₀)^(−1/2) 给出 Δd/d≈−Δn/n（近法向）。",
    )
    add_para(
        doc,
        "双角度一致性定义",
    )
    add_formula(doc, "ε = |d₁₀ − d₁₅| / [(d₁₀ + d₁₅)/2] × 100%", "10")
    add_follow(
        doc,
        "并以重采样标准差的逆方差加权得到联合厚度。式(10)检验的是“同一物理 d”假设，"
        "与单角度内部的峰谷一致性共同构成可靠性证据。",
    )

    add_para(doc, "6.3 结果分析", level="h2")
    add_table(
        doc,
        ["入射角", "峰序列/µm", "谷序列/µm", "双光束采用值/µm", "条件95%区间/µm"],
        [
            ["10°", "7.839", "7.926", "7.883", "[7.677, 8.152]"],
            ["15°", "7.811", "7.772", "7.791", "[7.560, 8.051]"],
        ],
        "表2  碳化硅双光束厚度反演结果（与 thickness_summary.csv 一致）",
    )
    add_follow(
        doc,
        "由表2：峰谷接近，说明极值识别稳定；两角度采用值相对差约 1.16%（式(10)），"
        "逆方差加权联合厚度",
    )
    add_formula(doc, "d_SiC ≈ 7.83 µm", "11")
    add_follow(
        doc,
        "式(11)即问题二主结论。统计区间反映噪声与局部条纹起伏；折射率 1% 的系统偏差约带来 1% 厚度反向偏差，"
        "不能被 bootstrap 消除。",
    )
    add_figure_slot(
        doc,
        "4",
        "双入射角厚度一致性（含 SiC/Si）",
        "output/angle_consistency.png",
        "SiC 两角度点应接近水平联合线",
    )
    add_figure_slot(
        doc,
        "5",
        "厚度对比与不确定度（常折射率结果与色散情景对照）",
        "output/thickness_comparison.png",
        "区分统计区间与掺杂系统范围",
    )

    add_para(doc, "6.4 小结", level="h2")
    add_para(
        doc,
        "问题二把式(9)落实为可运行的稳健测厚流程，得到 SiC 外延层厚度约 7.83 µm，双角度相对差约 1.16%。"
        "该结果将作为问题三多光束门控的对照基线：若证据不足，则不作 Airy 修正。",
    )

    # ========== 七、问题三 ==========
    add_para(doc, "七、问题三模型的建立与求解", level="h1")
    add_para(doc, "7.1 从双光束到多光束：公式如何扩展", level="h2")
    add_para(
        doc,
        "双光束只保留前两束反射。若计入外延层内第 j=1,2,… 次往返，则相邻往返之间多乘复因子",
    )
    add_formula(doc, "q = r₁₀ r₁₂ exp(−2αd/cosθ₁) exp(iδ)", "12")
    add_follow(
        doc,
        "式(12)中：r₁₀、r₁₂ 为界面振幅反射系数，指数吸收项抑制高次往返，exp(iδ) 则继承问题一的相位式(4)。"
        "当 |q| 不可忽略时，总反射场为等比级数求和，对 s/p 偏振有",
    )
    add_formula(doc, "r_tot = (r₀₁ + r₁₂ e^{2iβ}) / (1 + r₀₁ r₁₂ e^{2iβ})，  β = 2π d ν̃ n cosθ₁", "13")
    add_follow(
        doc,
        "式(13)的相位因子 β 与式(4)中的 δ/2 同源：都来自膜内往返光程。"
        "非偏振反射率取 |r^(s)|² 与 |r^(p)|² 的平均。对称无吸收时化为 Airy 形式",
    )
    add_formula(doc, "T = 1 / [1 + F sin²(δ/2)]，  F = 4Rᵢ/(1−Rᵢ)²，  R_Airy = 1 − T", "14")
    add_follow(
        doc,
        "式(14)表明：Rᵢ→0 时 F→0，R_Airy 回到缓慢变化背景，多光束退化为接近双光束的弱调制；"
        "Rᵢ 增大则条纹变尖，频谱中出现更强的高次谐波。这为后文用 A₂/A₁ 与 Rᵢ 做证据提供了机理。",
    )

    add_para(doc, "7.2 可观测条件与四项证据门控", level="h2")
    add_para(
        doc,
        "仅“理论上存在多次反射”不够，还需 |q|、相干性、分辨率与平行度达到可观察水平。"
        "数值上，本项目要求同时满足：",
    )
    add_formula(doc, "A₂/A₁ ≥ 0.08，  Rᵢ ≥ 0.12，  RMSE改善 ≥ 2%，  ΔAICc ≥ 10", "15")
    add_follow(
        doc,
        "其中基频由光学厚度预测：f₁ = 2 n cosθ₁ · d · 10⁻⁴，谐波比在 2f₁ 邻域取幅值比。"
        "Rᵢ 可被未建模基线部分吸收，故不能单独作为证据；式(15)的“且”关系用于防止过拟合误判。",
    )
    add_figure_slot(
        doc,
        "6",
        "多光束四项证据矩阵（PASS/FAIL）",
        "output/multibeam_evidence.png",
        "Si 应多为通过，SiC 谐波/反射率应显示不足",
    )
    add_figure_slot(
        doc,
        "7",
        "双光束与 Airy 模型质量对比（RMSE 及改善率）",
        "output/model_quality.png",
        "结合表3解读：残差改善不等于可观测多光束",
    )

    add_para(doc, "7.3 结果分析：硅片与碳化硅的分叉结论", level="h2")
    add_table(
        doc,
        ["数据集", "A₂/A₁", "有效反射率 Rᵢ", "RMSE改善/%", "ΔAICc", "判定"],
        [
            ["SiC 10°", "0.019", "0.002", "23.2", "3072", "证据不足"],
            ["SiC 15°", "0.037", "0.005", "25.9", "3480", "证据不足"],
            ["Si 10°", "0.196", "0.410", "11.9", "1475", "可观测多光束"],
            ["Si 15°", "0.222", "0.422", "8.65", "1051", "可观测多光束"],
        ],
        "表3  多光束四项证据诊断（与程序输出一致）",
    )
    add_follow(
        doc,
        "表3显示：Si 的谐波比约 0.20–0.22，Rᵢ约 0.41–0.42，四项同时满足式(15)；"
        "SiC 虽有 RMSE/AICc 改善，但谐波比仅 0.02–0.04，Rᵢ接近下界，故判定不可观测，主结果不修正。",
    )
    add_figure_slot(
        doc,
        "8",
        "二次谐波与基频幅值比（原始证据图）",
        "output/raw_evidence/multibeam/harmonic_ratio_raw.png",
        "Si 明显高于阈值，SiC 远低于阈值",
    )
    add_figure_slot(
        doc,
        "9",
        "有效反射率原始证据图",
        "output/raw_evidence/multibeam/effective_reflectivity_raw.png",
    )
    add_figure_slot(
        doc,
        "10",
        "附件3（Si，10°）谐波频谱",
        "output/raw_evidence/multibeam/si_10_harmonic_spectrum_raw.png",
        "应能看到基频主峰；二次谐波相对较强",
    )
    add_figure_slot(
        doc,
        "11",
        "附件1（SiC，10°）谐波频谱",
        "output/raw_evidence/multibeam/sic_10_harmonic_spectrum_raw.png",
        "与图10对比：二次谐波很弱",
    )
    add_figure_slot(
        doc,
        "12",
        "附件3（Si，10°）完整拟合证据图",
        "output/si_10_fit.png",
    )
    add_figure_slot(
        doc,
        "13",
        "附件4（Si，15°）完整拟合证据图",
        "output/si_15_fit.png",
    )

    add_table(
        doc,
        ["对象", "双光束/µm", "Airy多光束/µm", "最终模型", "采用厚度/µm"],
        [
            ["Si 10°", "3.403", "3.598", "多光束", "3.598"],
            ["Si 15°", "3.410", "3.584", "多光束", "3.584"],
            ["SiC 10°", "7.883", "8.449", "双光束", "7.883"],
            ["SiC 15°", "7.791", "8.333", "双光束", "7.791"],
        ],
        "表4  厚度对比与最终采用值",
    )
    add_follow(doc, "硅片两角度多光束厚度相对差约 0.38%，加权联合")
    add_formula(doc, "d_Si ≈ 3.59 µm", "16")
    add_follow(
        doc,
        "若硅片仍用双光束，联合厚度约 3.41 µm，相对式(16)偏低约 5%，说明忽略多次反射会带来模型误差。"
        "对 SiC，Airy 厚度系统性偏高，但因不满足式(15)，仍采用问题二的双光束结果式(11)。",
    )

    add_para(doc, "7.4 小结", level="h2")
    add_para(
        doc,
        "问题三在问题一相位结构上引入多光束级数求和（式(12)–(14)），并用式(15)做可观测门控。"
        "结论：附件3、4存在可观察多光束干涉，d_Si≈3.59 µm；附件1、2证据不足，SiC 不作多光束修正，d_SiC≈7.83 µm。",
    )

    # ========== 八、汇总 ==========
    add_para(doc, "八、问题一二三结果汇总", level="h1")
    add_table(
        doc,
        ["问题", "核心公式链", "采用策略", "关键结果"],
        [
            ["一", "(1)→(9)", "斜入射双光束解析", "d=1/[2Δν̃√(n²−sin²θ₀)]"],
            ["二", "(9)+(10)", "透明窗+峰谷稳健+双角度融合", "d_SiC≈7.83 µm"],
            ["三", "(12)→(16)", "Airy拟合+四项证据门控", "d_Si≈3.59 µm；SiC不修正"],
        ],
        "表5  三问公式链与结论对照",
    )
    add_para(
        doc,
        "插图文件均来自本项目 output/ 目录；请按文中【插图预留】标注逐一插入。"
        "数值以 output/thickness_summary.csv 与 consistency.json 为准。",
        first_indent=False,
    )

    return doc


def write_md():
    MD_OUT.write_text(
        """# 2025B 论文详细版（问题一二三）

## 公式主链
1. Snell → cosθ₁
2. 光程差 ΔL=2nd cosθ₁
3. 相位差 δ=4π d ν̃ n cosθ₁+φᵣ
4. 相位坐标 g=ν̃√(n²−sin²θ₀)
5. 厚度 d=1/[2Δν̃√(n²−sin²θ₀)]
6. 多光束：q 因子 → Airy → 四项证据门控

## 关键数值
- SiC：7.883 / 7.791 µm，相对差 1.16%，联合 **7.83 µm**（双光束）
- Si：3.598 / 3.584 µm，相对差 0.38%，联合 **3.59 µm**（多光束）

## 插图清单（按文中图号）
1. model_flowchart.png
2. sic_10_fit.png
3. sic_15_fit.png
4. angle_consistency.png
5. thickness_comparison.png
6. multibeam_evidence.png
7. model_quality.png
8. raw_evidence/multibeam/harmonic_ratio_raw.png
9. raw_evidence/multibeam/effective_reflectivity_raw.png
10. raw_evidence/multibeam/si_10_harmonic_spectrum_raw.png
11. raw_evidence/multibeam/sic_10_harmonic_spectrum_raw.png
12. si_10_fit.png
13. si_15_fit.png

详见 Word：`2025B论文_问题一二三_详细版.docx`
""",
        encoding="utf-8",
    )


def main():
    doc = build()
    doc.save(str(OUT))
    write_md()
    # 尝试更新常用副本名
    alt = DOC_DIR / "2025B论文_问题一二三.docx"
    try:
        doc.save(str(alt))
        alt_status = "updated"
    except PermissionError:
        alt_status = "locked"
    print("OUT", OUT.as_posix())
    print("ALT", alt_status, alt.as_posix())
    print("MD", MD_OUT.as_posix())


if __name__ == "__main__":
    main()
