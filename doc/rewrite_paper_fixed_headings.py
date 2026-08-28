# -*- coding: utf-8 -*-
"""按既定一/二级标题重写详细版：不改标题名，模型建立重点展开。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_DIR = Path(__file__).resolve().parent
OUT = DOC_DIR / "2025B论文_问题一二三_详细版.docx"
OUT_ALT = DOC_DIR / "2025B论文_问题一二三_详细版_v7方案B.docx"
MD_OUT = DOC_DIR / "2025B论文_问题一二三_详细版.md"


def set_run(run, *, east="宋体", size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east)


def add(doc, text, *, kind="body", indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74) if (kind == "body" and indent) else Pt(0)

    if kind == "title":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), east="黑体", size=15, bold=True)
    elif kind == "h1":
        pf.space_before = Pt(12)
        set_run(p.add_run(text), east="黑体", size=15, bold=True)
    elif kind == "h2":
        pf.space_before = Pt(8)
        set_run(p.add_run(text), east="黑体", size=12, bold=True)
    elif kind == "h3":
        set_run(p.add_run(text), east="黑体", size=10.5, bold=True)
    elif kind == "caption":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), east="黑体", size=10.5, bold=True)
    elif kind == "formula":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), east="宋体", size=10.5)
    elif kind == "slot":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), east="宋体", size=10.5, color=RGBColor(0x80, 0x00, 0x00))
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "12")
            el.set(qn("w:space"), "4")
            el.set(qn("w:color"), "C00000")
            pBdr.append(el)
        pPr.append(pBdr)
    else:
        set_run(p.add_run(text), east="宋体", size=10.5)
    return p


def formula(doc, text, no):
    add(doc, f"{text}　　({no})", kind="formula")


def follow(doc, text):
    add(doc, text, indent=False)


def slot(doc, no, title, path, tip=""):
    msg = f"【插图预留：图{no}】\n请插入：{path}\n建议居中，宽约 12–14 cm"
    if tip:
        msg += f"\n读图要点：{tip}"
    add(doc, msg, kind="slot")
    add(doc, f"图{no}  {title}", kind="caption")


def table(doc, headers, rows, caption):
    add(doc, caption, kind="caption")
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(h), east="黑体", size=10.5, bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(v), east="宋体", size=10.5)
    doc.add_paragraph()


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    add(doc, "碳化硅外延层厚度的红外干涉测量及多光束干涉影响分析", kind="title")

    # ===== 既定一级/二级标题：问题重述 =====
    add(doc, "问题重述", kind="h1")
    add(doc, "1.1 问题背景", kind="h2")
    add(
        doc,
        "碳化硅作为典型的第三代半导体材料，凭借宽禁带、高热导率、高击穿电场等突出综合性能，在功率电子、射频器件等领域得到广泛研究与工程应用。"
        "碳化硅外延层厚度是评价外延材料品质的核心指标，其数值精度会直接影响半导体器件的击穿电压、导通电阻等关键性能参数。"
        "因此，建立一套科学、精准且具备可靠性的碳化硅外延层厚度检测标准，对碳化硅器件工艺研发、质量管控具有十分重要的现实意义。"
        "图1为多尺寸SiC晶圆示意图。",
    )
    add(doc, "图1  多尺寸SiC晶圆", kind="caption")
    add(
        doc,
        "红外干涉法属于非破坏性检测手段，适用于外延层厚度测试。该方法的基本原理为：碳化硅外延层与衬底的载流子掺杂浓度存在差异，使得二者折射率不同；"
        "当红外光照射样品时，一部分光束由外延层上表面直接反射，另一部分光束穿透外延层后在衬底界面发生反射，两束反射光相互叠加形成干涉条纹，"
        "结合红外光谱波长、外延层折射率、光线入射角等参数，即可计算得到外延层厚度。",
    )
    add(
        doc,
        "本文针对红外干涉法测定碳化硅以及硅单体外延层厚度这一情景，在不同反射率光谱、入射角以及外延层–衬底双界面干涉机制已经给定的条件下，"
        "综合考虑透明波段选取、多光束物理模型、折射率随波长与载流子浓度的色散，以及对不同角度测量结果的一致性约束，"
        "建立可解释的测厚与诊断流程，以得到可靠的外延层厚度估计，并正确判定是否需要进行多光束修正。",
    )

    add(doc, "1.2 问题提出", kind="h2")
    add(
        doc,
        "外延层厚度直接影响功率器件性能，是工艺控制的关键指标。红外光在空气–外延层与外延层–衬底两界面反射后发生干涉，反射率随波数出现条纹；"
        "在适当简化下，条纹周期与膜厚有明确关系，因而红外干涉法具有非接触、无损的优点。但实际光谱中，强吸收带与过渡区常使常数折射率假设失效，"
        "噪声与基线会影响极值识别，多次反射还可能使条纹偏离简单余弦。波段或模型选择不当，将引入系统偏差。据此，本文分三个层次展开。",
    )
    add(
        doc,
        "问题一：要求在双光束假设下建立斜入射薄膜干涉厚度模型：由 Snell 定律确定膜内折射角，写出往返光程差与相位差，"
        "得到常折射率下厚度与相邻同类极值波数间隔的解析关系，为后续反演提供依据。",
    )
    add(
        doc,
        "问题二：要求利用附件一、二的碳化硅光谱，确定同一块晶圆的外延层厚度。两份数据分别对应十度与十五度入射角，"
        "需兼顾透明波段选取与条纹去噪，并以同一物理厚度检验双角度一致性，同时评价结果可靠性。",
    )
    add(
        doc,
        "问题三：进一步讨论多光束干涉：先明确其可观测条件，再判断附件三、四硅片光谱是否存在显著多光束特征并建模测厚；"
        "同时检验附件一、二的碳化硅结果是否需要多光束修正，避免仅凭残差下降而误判。",
    )
    add(doc, "图2  问题关系图", kind="caption")

    # ===== 问题分析（标题原样保留）=====
    add(doc, "问题分析", kind="h1")
    add(doc, "2.2问题一分析", kind="h2")
    add(
        doc,
        "问题一的关键是把斜入射几何正确写入相位差，得到可供后续反演直接使用的解析关系。"
        "由折射定律确定膜内折射角后，写出一次往返光程差与相位差；相邻同类极值对应相位改变一个完整周期，从而得到厚度与波数间隔的对应关系。"
        "该关系在入射角趋于零时应退化为常见的正入射公式，可作为推导正确性的退化检验。"
        "问题一不直接给出数值厚度，而是为问题二、三提供统一的理论锚点。",
    )
    add(doc, "2.2问题二分析", kind="h2")
    add(
        doc,
        "问题二在问题一结论基础上，利用附件一、二确定同一块碳化硅晶圆的外延层厚度并评价可靠性。"
        "难点在于全谱并非处处满足常折射率透明假设：强吸收带与过渡区会扭曲条纹，使频域估计与极值估计结果分裂。"
        "因此先对多个候选透明窗评分，再在选定窗内做双尺度去基线，以频域变换粗估厚度、以峰谷稳健回归为主估计，并结合有界精修与重抽样给出区间。"
        "最后以同一物理厚度约束检验十度与十五度结果的一致性，完成双角度融合。",
    )
    add(doc, "2.2问题三分析", kind="h2")
    add(
        doc,
        "问题三在问题二所用透明窗与预处理基础上，讨论是否需要从双光束升级到多光束。"
        "难点不在于能否把残差拟合得更小，而在于多光束是否真正可观测。"
        "为此先给出多光束可观测的必要条件，再建立多次反射干涉模型，并用谐波比、有效反射率、信息准则与误差改善等指标联合诊断。"
        "对附件三、四的硅片，若诊断为多光束显著则采用多光束模型测厚；对附件一、二的碳化硅，若证据不足则不作多光束修正，保留问题二结论。"
        "三问层层递进：问题一给公式，问题二完成碳化硅主反演，问题三作模型分叉判定。",
    )

    # ===== 符号说明 =====
    add(doc, "符号说明", kind="h1")
    add(doc, "本文主要符号见表1，正文公式与该表保持一致。")
    table(
        doc,
        ["符号", "含义", "单位"],
        [
            ["θ₀", "空气中入射角", "°"],
            ["θ₁", "外延层内折射角", "°"],
            ["n", "外延层折射率", "—"],
            ["ν̃", "波数", "cm⁻¹"],
            ["d", "外延层厚度", "µm"],
            ["ΔL", "一次往返光程差", "cm"],
            ["δ", "相位差", "rad"],
            ["φᵣ", "界面反射附加相位", "rad"],
            ["g", "光学相位坐标", "cm⁻¹"],
            ["Δν̃", "相邻同类极值波数间隔", "cm⁻¹"],
            ["A,B", "慢变基线与条纹振幅", "%"],
            ["Rᵢ", "Airy有效界面反射率", "—"],
            ["A₂/A₁", "二次谐波与基频幅值比", "—"],
            ["ΔAICc", "双光束相对多光束AICc差", "—"],
        ],
        "表1  主要符号说明",
    )

    # ===== 问题一 =====
    add(doc, "问题一", kind="h1")

    add(doc, "4.1数据预处理", kind="h2")
    add(
        doc,
        "问题一以机理建模为主，但仍需明确后续数值实现所依赖的数据约定。"
        "四份附件各约 7469 点，波数约 399.7–4000.1 cm⁻¹，间隔约 0.482 cm⁻¹。"
        "读取时进行有限值过滤、波数升序与去重，并剔除首点可疑零值。"
        "问题一本身不输出厚度数值；预处理规则将直接服务于问题二、三。",
    )

    add(doc, "4.2模型建立", kind="h2")
    add(
        doc,
        "本节建立斜入射双光束干涉厚度模型。建模按物理含义分四步展开：先确定膜内折射角，再写出一次往返光程与相位，"
        "进而给出双光束反射率的振荡形式，最后由同类极值波数间隔反演厚度。"
        "各步给出主要公式与物理含义；公式之间只作简要承接，不展开代数推导细节，最终得到可供问题二直接调用的解析关系。",
    )

    add(doc, "4.2.1 折射角确定", kind="h3")
    add(
        doc,
        "斜入射时，空气侧入射角 θ₀ 可测，但干涉相位由膜内传播方向决定。"
        "设外延层折射率为 n、膜内折射角为 θ₁。由 Snell 定律得到",
    )
    formula(doc, "sin θ₀ = n sin θ₁", "1")
    follow(
        doc,
        "该关系把外部几何条件映射到膜内光线方向，是后续光程与相位修正的起点。"
        "在实折射率且未发生全反射的条件下，进一步得到",
    )
    formula(doc, "cos θ₁ = √(1 − sin²θ₀ / n²)", "2")
    follow(
        doc,
        "上式要求 n>sinθ₀；本文附件角度与题面常折射率均满足该条件。"
        "cos θ₁ 将进入光程差与相位表达式，是斜入射相对正入射的核心修正因子。",
    )

    add(doc, "4.2.2 光程差与相位差", kind="h3")
    add(
        doc,
        "双光束假设下，一束光在外延层上表面直接反射；另一束进入膜内，经下界面反射后再射出，两束在探测器处相干叠加。"
        "膜内一次往返对应的光程差可写为",
    )
    formula(doc, "ΔL = 2 n d cos θ₁", "3")
    follow(
        doc,
        "由此可知，厚度 d 通过光学因子 n cosθ₁ 进入可观测量。"
        "由光程差与波数关系，并计及界面反射附加相位 φᵣ，得到相位差",
    )
    formula(doc, "δ(ν̃) = (2π/λ)ΔL + φᵣ = 4π d ν̃ n cos θ₁ + φᵣ", "4")
    follow(
        doc,
        "上式说明相位随波数近似线性增长；其中 φᵣ 主要影响极值对应的干涉级次奇偶，"
        "一般不改变相邻同类极值的波数间隔，故厚度反演可主要依赖间隔信息。"
        "为书写紧凑，引入光学相位坐标",
    )
    formula(doc, "g(ν̃,θ₀) = ν̃ n cos θ₁ = ν̃ √(n² − sin²θ₀)", "5")
    follow(
        doc,
        "结合折射角余弦与相位差表达式，相位亦可写为 δ=4πdg+φᵣ。"
        "于是，反射谱在波数轴上的干涉周期主要由 g 的变化率决定。",
    )

    add(doc, "4.2.3 双光束反射率形式", kind="h3")
    add(
        doc,
        "两束反射场叠加后，仪器基线与条纹振幅通常随波数缓慢变化，分别记为 A(ν̃)、B(ν̃)。"
        "双光束反射率可写为",
    )
    formula(doc, "R(ν̃) = A(ν̃) + B(ν̃) cos δ(ν̃)", "6")
    follow(
        doc,
        "上式把观测光谱分解为慢变包络与快变干涉项：快变部分由相位 δ 控制，慢变部分吸收基线与振幅漂移。"
        "这为数值处理提供直接依据——先去基线抑制 A，再对残差做频域或峰谷分析，以提取与厚度相关的振荡信息。",
    )

    add(doc, "4.2.4 厚度与波数间隔关系", kind="h3")
    add(
        doc,
        "设两同类极值 ν̃_a、ν̃_b 之间跨越 Δm 个干涉级次，则其间相位改变 2πΔm。"
        "由相位表达式与光学坐标定义，得到",
    )
    formula(doc, "4π d [g(ν̃_b) − g(ν̃_a)] = 2π Δm", "7")
    follow(doc, "由此得到厚度的一般表达式")
    formula(doc, "d = Δm / {2 [g(ν̃_b,θ₀) − g(ν̃_a,θ₀)]}", "8")
    follow(
        doc,
        "在窄波段内折射率近似为常数，并取相邻同类极值（Δm=1）时，由上式得到常用测厚公式",
    )
    formula(doc, "d = 1 / [2 Δν̃ √(n² − sin²θ₀)]", "9")
    follow(
        doc,
        "当 d 以 µm、Δν̃ 以 cm⁻¹ 计时，上式右侧再乘 10⁴。"
        "正入射极限下，该测厚公式退化为 d=1/(2nΔν̃)，与经典正入射公式一致，可作为斜入射推广正确性的退化检验。"
        "该结果即问题一的核心结论，也是问题二厚度反演的直接入口。",
    )

    add(doc, "4,3 模型求解", kind="h2")
    add(
        doc,
        "问题一以解析推导为主。对应的数值求解接口为：由光谱估计 Δν̃，再代入测厚公式计算厚度。"
        "具体估计 Δν̃ 的流程在问题二实现，包括 FFT 粗估与峰谷 Theil–Sen 稳健回归。",
    )

    add(doc, "4.4 结果分析", kind="h2")
    add(
        doc,
        "问题一给出可检验的解析映射而非数值厚度。其作用有三：把斜入射写入相位；说明厚度由同类极值间隔决定；"
        "明确常折射率假设的适用范围，为问题二选带与问题三模型分叉提供边界。",
    )
    slot(doc, "3", "数学模型流程图", "output/model_flowchart.png", "对应问题一至三总体流程")

    add(doc, "4.5 总结", kind="h2")
    add(
        doc,
        "本节完成双光束斜入射厚度模型，得到厚度与波数间隔的解析关系。后续问题二将在透明波段内估计 Δν̃ 并计算碳化硅厚度；"
        "问题三将在同一相位结构上扩展为多光束模型。",
    )

    # ===== 问题二 =====
    add(doc, "问题二", kind="h1")

    add(doc, "5.1数据预处理", kind="h2")
    add(
        doc,
        "附件一、二分别对应同一 SiC 晶圆在 10°、15° 的反射谱。"
        "为使常折射率测厚公式成立，需避开强声子吸收带。本项目对候选透明窗评分后，主拟合波段约取 1200–4000 cm⁻¹。"
        "短窗 Savitzky–Golay 去噪，长窗提取慢变基线，残差近似对应双光束反射率中的振荡项。"
        "附件二最大反射率约 102.74%，提示绝对标定可能异常；双光束厚度主要由峰位/间距决定，故保留原始量程并记入审计。",
    )
    slot(doc, "4", "附件1（SiC，10°）处理与拟合证据图", "output/sic_10_fit.png")
    slot(doc, "5", "附件2（SiC，15°）处理与拟合证据图", "output/sic_15_fit.png")

    add(doc, "5.2模型建立", kind="h2")
    add(
        doc,
        "问题二在厚度–波数间隔关系之上建立可计算的测厚模型：把“如何得到 Δν̃”与“如何融合两角度、如何量化不确定度”明确化。"
        "模型由条纹测厚子模型与可靠性评价子模型组成。",
    )

    add(doc, "5.2.1 条纹测厚子模型", kind="h3")
    add(
        doc,
        "对预处理后的残差谱，厚度估计分三层组织，目标都是稳定得到测厚所需的 Δν̃（或等价地得到 d）。"
        "第一层为频域粗估：残差主频与厚度近似满足",
    )
    formula(doc, "f₁ ≈ 2 n cosθ₁ · d · 10⁻⁴", "10")
    follow(
        doc,
        "由相位振荡周期关系可得到上式。它用于给出厚度初值，并确定峰/谷检测的尺度参数。"
        "第二层为峰谷稳健估计：分别提取峰、谷序列，对序号–波数做 Theil–Sen 回归，斜率即为稳健 Δν̃；"
        "再代入测厚公式得到 d_peak、d_valley，并综合为双光束采用值 d_two。"
        "第三层为有界精修：在稳健解邻域内优化相位模型；若宽波段色散导致级次跳变，则保留峰谷解以避免系统偏差。"
        "折射率取 n_SiC=2.55。",
    )

    add(doc, "5.2.2 可靠性评价子模型", kind="h3")
    add(
        doc,
        "可靠性从统计不确定度与系统误差两方面评价。"
        "统计不确定度由峰/谷间距重采样得到条件 95% 区间；系统误差主要来自折射率假设。"
        "由测厚公式可知厚度对折射率的依赖关系为",
    )
    formula(doc, "d ∝ (n² − sin²θ₀)^(−1/2)", "11")
    follow(
        doc,
        "近法向时，由上式可得到相对误差近似 Δd/d≈−Δn/n：例如 n 偏差 1%，厚度约反向偏差 1%，且无法被重采样消除。"
        "同一晶圆应共享物理厚度，据此定义双角度相对差",
    )
    formula(doc, "ε = |d₁₀ − d₁₅| / [(d₁₀+d₁₅)/2] × 100%", "12")
    follow(
        doc,
        "并以重采样标准差的逆方差加权得到联合厚度。"
        "该相对差与峰谷内部一致性共同构成可靠性判据：相对差较小且峰谷接近时，认为双光束反演可信。",
    )

    add(doc, "5.2.3 本征色散情景模型", kind="h3")
    add(
        doc,
        "为分析常折射率假设带来的系统误差，本文增加本征色散情景模型。"
        "硅采用 Edwards–Ochoa 红外本征色散关系；碳化硅采用固定声子参数的单振子晶格模型。"
        "在本征情景中关闭自由载流子项，在低、中、高掺杂情景中使用预先给定的固定浓度，载流子浓度不作为优化变量。"
        "材料的复折射率统一写为",
    )
    formula(doc, "ñ(ν̃)=√[ε_lattice(ν̃)+ε_Drude(ν̃,N_fixed)]", "13")
    follow(
        doc,
        "其中 N_fixed 由情景预先指定。本征情景取 N_fixed=0。"
        "由于外延层与衬底采用相同本征介电函数时界面反射会退化，本模型不利用绝对反射幅值反演厚度，"
        "而在透明区用复折射率实部修正干涉相位，并通过正余弦变量投影吸收未知振幅、基线与相位偏置。"
        "双角度共享的唯一关键非线性变量为厚度。",
    )
    formula(doc, "d_s=arg min_d Σ_a ‖y_a−X_a[d,n_s(ν̃)]β̂_a‖²", "14")
    follow(
        doc,
        "依次计算本征、低、中、高四种固定情景的厚度，并将全部情景结果与常折射率主值共同形成系统误差包络。"
        "该包络只用于可靠性分析，不按拟合误差自动替换主厚度。",
    )

    add(doc, "5.2.4 自由浓度审计", kind="h3")
    add(
        doc,
        "原有载流子复介电联合反演仍完整保留，但调整为审计路径。"
        "该路径同时开放厚度、外延层浓度与衬底浓度，并检查 Jacobian 条件数、参数相关性、先验边界、"
        "连续留段厚度稳定性、固定情景比较及浓度轮廓区间。"
        "只有全部证据支持时才允许解释浓度点估计；否则候选值仅作为不可辨识证据，不进入论文主结论。",
    )

    add(doc, "5.2.5 色散坐标多峰谷模型", kind="h3")
    add(
        doc,
        "为同时保留极值法对基线的稳健性和色散模型对折射率变化的描述，本文进一步建立色散坐标多峰谷共享厚度模型。"
        "模型仍以透明波段筛选为前置条件，在本征及固定掺杂情景下，先将峰谷波数映射为光学相位坐标",
    )
    formula(doc, "g_s,a(ν̃)=ν̃√[n_s²(ν̃)−sin²θ_a]", "15")
    follow(
        doc,
        "其中下标 s 表示本征、低、中、高固定折射率情景，下标 a 表示入射角。"
        "在该坐标中，同类极值的干涉级次近似线性增加。"
        "对10°峰、10°谷、15°峰和15°谷四条序列分别设置独立截距，以吸收未知绝对级次和峰谷反射相位差；四条序列共享同一厚度斜率：",
    )
    formula(doc, "m_i=2d_cm g_s,a(ν̃_i)+c_j", "16")
    follow(
        doc,
        "相邻光学坐标间距相对中位间距的整数倍用于识别漏峰或漏谷；"
        "共享斜率采用稳健损失估计，并按残差中位绝对偏差审计异常极值。"
        "模型分别计算峰序列、谷序列、两个角度和连续三波段的厚度，只有这些结果同时稳定时，才采用本征色散情景作为名义厚度。"
        "低、中、高固定情景与常折射率结果共同构成模型系统范围。",
    )

    add(doc, "5.3 模型求解", kind="h2")
    add(
        doc,
        "对附件一、二分别执行：选带 → 预处理 → FFT粗估 → 峰谷Theil–Sen → 有界精修 → 间距bootstrap；"
        "再计算双角度相对差并做逆方差加权融合。"
        "随后执行本征与三档固定掺杂情景，只优化共享厚度并形成系统误差范围；"
        "在此基础上提取多组峰谷并映射到色散坐标，对四条序列进行共享厚度回归，"
        "再执行峰谷、角度、留段与重采样检验。最后运行保留的自由浓度审计路径，输出失败指标与回退依据。"
        "全过程使用同一阈值与随机种子规则，保证两角度可比。",
    )

    add(doc, "5.4 结果分析", kind="h2")
    table(
        doc,
        ["入射角", "峰序列/µm", "谷序列/µm", "双光束采用值/µm", "条件95%区间/µm"],
        [
            ["10°", "7.839", "7.926", "7.883", "[7.677, 8.152]"],
            ["15°", "7.811", "7.772", "7.791", "[7.560, 8.051]"],
        ],
        "表2  碳化硅双光束厚度结果",
    )
    follow(
        doc,
        "由表2，峰谷接近；两角度相对差约 1.16%。加权联合厚度为",
    )
    formula(doc, "d_SiC,const ≈ 7.83 µm", "17")
    add(
        doc,
        "常折射率结果作为基线。升级后的色散坐标多峰谷模型得到本征情景名义厚度",
    )
    formula(doc, "d_SiC,v8 ≈ 7.452 µm", "18")
    follow(
        doc,
        "其条件统计95%区间约为[7.431,7.483] µm。峰、谷子模型分别约为7.434 µm和7.463 µm，相对差0.399%；"
        "10°和15°结果分别约为7.500 µm和7.412 µm，相对差1.175%；连续留段CV为0.738%，最大偏移为1.563%。"
        "39个合格极值均被保留，全部稳定性条件通过，因此采用7.452 µm作为SiC名义厚度。"
        "低、中、高固定情景及常折射率基线共同形成约7.368–7.832 µm的模型系统范围。"
        "该范围描述折射率设定的不确定性，不是统计置信区间，也不表示测得了载流子浓度。"
        "自由浓度审计中，连续留段厚度变异系数约 8.69%，最大偏移约 14.99%，均超过稳定性门槛；"
        "自由拟合误差还劣于最佳固定情景，因此浓度不可唯一辨识。",
    )
    slot(doc, "6", "双入射角厚度一致性", "output/angle_consistency.png", "SiC 两点应接近联合水平线")
    slot(doc, "7", "厚度对比与不确定度", "output/thickness_comparison.png")
    slot(doc, "8", "SiC厚度反演波段与极值资格", "output/analysis_evidence/sic_band_eligibility_analysis.png")
    slot(doc, "9", "SiC色散坐标多峰谷共享厚度回归", "output/analysis_evidence/sic_order_fit_analysis.png")
    slot(doc, "10", "SiC稳定性与模型采用门控", "output/analysis_evidence/sic_stability_gates_analysis.png")

    add(doc, "5.5 总结", kind="h2")
    add(
        doc,
        "问题二在常折射率条件下得到约7.83 µm的基线结果；升级后的色散坐标多峰谷模型得到SiC名义厚度约7.45 µm，"
        "条件统计95%区间约为[7.43,7.48] µm，固定色散情景与常折射率基线给出约7.37–7.83 µm的模型系统范围。"
        "自由浓度审计证明现有数据不能支持唯一浓度点估计，因此增强结果仅用于可靠性与模型评价。"
        "该结果作为问题三多光束判定的对照基线。",
    )

    # ===== 问题三 =====
    add(doc, "问题三", kind="h1")

    add(doc, "6.1数据预处理", kind="h2")
    add(
        doc,
        "附件三、四为硅片 10°、15° 光谱；附件一、二继续用于检验碳化硅是否需要多光束修正。"
        "预处理与问题二一致：选带、双尺度滤波与去基线。"
        "多光束诊断在残差谱上进行，同时拟合 Airy 模型以便与双光束比较 RMSE 与 AICc。",
    )

    add(doc, "6.2模型建立", kind="h2")
    add(
        doc,
        "问题三模型包含三块：多光束可观测条件、Airy/Fresnel 反射模型、四项证据门控。"
        "它们建立在问题一相位结构之上，并决定是否从双光束升级。",
    )

    add(doc, "6.2.1 多光束可观测条件", kind="h3")
    add(
        doc,
        "当界面反射足够强且吸收不太强时，除双光束外还存在更高次往返。"
        "计入第 j 次往返后，相邻往返之间可多乘复因子",
    )
    formula(doc, "q = r₁₀ r₁₂ exp(−2αd/cosθ₁) exp(iδ)", "19")
    follow(
        doc,
        "其中，指数吸收项抑制高次往返，相位因子与问题一中的 δ 一致。"
        "多光束要在光谱上可观测，要求 |q| 足够大，并同时满足相干性、仪器分辨率与膜厚均匀等条件。"
        "因此，不能仅凭“物理上存在多次反射”就直接改用多光束模型，而需结合可观测证据判定。",
    )

    add(doc, "6.2.2 Airy反射模型", kind="h3")
    add(
        doc,
        "在多次反射可观测时，采用单层薄膜的全反射振幅求和。"
        "对 s/p 偏振，总反射振幅可写为",
    )
    formula(doc, "r_tot = (r₀₁ + r₁₂ e^{2iβ}) / (1 + r₀₁ r₁₂ e^{2iβ})， β=2π d ν̃ n cosθ₁", "20")
    follow(
        doc,
        "其中 β 与问题一相位结构同源。对称无吸收情形下，由上式得到 Airy 形式",
    )
    formula(doc, "T=1/[1+F sin²(δ/2)]， F=4Rᵢ/(1−Rᵢ)²， R_Airy=1−T", "21")
    follow(
        doc,
        "上式表明：有效界面反射率 Rᵢ 较小时精细度 F 很小，多光束调制弱，近似退回双光束；"
        "Rᵢ 增大时条纹变尖、高次谐波增强，这正是后续谐波比诊断的物理依据。"
        "数值实现中，非线性参数取 (d,Rᵢ,相位偏置)，线性增益与慢变基线采用剖面消元。",
    )

    add(doc, "6.2.3 四项证据门控", kind="h3")
    add(
        doc,
        "多参数模型拟合残差通常更容易下降，若仅凭 RMSE 改善，容易把“可拟合”误判为“多光束可观测”。"
        "为此建立四项证据门控，要求同时满足",
    )
    formula(doc, "A₂/A₁≥0.08， Rᵢ≥0.12， RMSE改善≥2%， ΔAICc≥10", "22")
    follow(
        doc,
        "其中基频由频域粗估关系预测，谐波比在 2f₁ 邻域取值；Rᵢ 可能被未建模基线部分吸收，故不能单独作为证据。"
        "门控条件以逻辑“且”连接四项：全部通过则采用多光束厚度，否则回退双光束厚度。",
    )

    add(doc, "6.3 模型求解", kind="h2")
    add(
        doc,
        "对四份光谱分别拟合双光束与 Airy 模型，计算谐波频谱与四项证据指标并判定；"
        "对判定为多光束的硅片两角度结果再做一致性检验与逆方差加权。",
    )

    add(doc, "6.4 结果分析", kind="h2")
    table(
        doc,
        ["数据集", "A₂/A₁", "Rᵢ", "RMSE改善/%", "ΔAICc", "判定"],
        [
            ["SiC 10°", "0.019", "0.002", "23.2", "3072", "证据不足"],
            ["SiC 15°", "0.037", "0.005", "25.9", "3480", "证据不足"],
            ["Si 10°", "0.196", "0.410", "11.9", "1475", "可观测多光束"],
            ["Si 15°", "0.222", "0.422", "8.65", "1051", "可观测多光束"],
        ],
        "表3  多光束四项证据诊断",
    )
    follow(
        doc,
        "表3表明：硅片满足四项证据门控；碳化硅虽有残差改善，但谐波比与 Rᵢ 不足，故不作多光束修正。",
    )
    slot(doc, "11", "多光束证据矩阵", "output/multibeam_evidence.png")
    slot(doc, "12", "模型质量对比", "output/model_quality.png")
    slot(doc, "13", "谐波比原始图", "output/raw_evidence/multibeam/harmonic_ratio_raw.png")
    slot(doc, "14", "有效反射率原始图", "output/raw_evidence/multibeam/effective_reflectivity_raw.png")
    slot(doc, "15", "Si 10°谐波频谱", "output/raw_evidence/multibeam/si_10_harmonic_spectrum_raw.png")
    slot(doc, "16", "SiC 10°谐波频谱", "output/raw_evidence/multibeam/sic_10_harmonic_spectrum_raw.png")
    slot(doc, "17", "附件3拟合证据图", "output/si_10_fit.png")
    slot(doc, "18", "附件4拟合证据图", "output/si_15_fit.png")
    slot(doc, "19", "Si色散坐标模型稳定性与回退依据", "output/analysis_evidence/si_stability_gates_analysis.png")
    slot(doc, "20", "最终厚度模型决策", "output/analysis_evidence/final_model_decision_analysis.png")

    table(
        doc,
        ["对象", "双光束/µm", "Airy/µm", "最终模型", "采用厚度/µm"],
        [
            ["Si 10°", "3.403", "3.598", "多光束", "3.598"],
            ["Si 15°", "3.410", "3.584", "多光束", "3.584"],
            ["SiC 10°", "7.883", "8.449", "双光束", "7.883"],
            ["SiC 15°", "7.791", "8.333", "双光束", "7.791"],
        ],
        "表4  厚度对比与最终采用值",
    )
    follow(doc, "硅片两角度相对差约 0.38%，加权联合")
    formula(doc, "d_Si ≈ 3.59 µm", "23")
    follow(
        doc,
        "若硅片仍用双光束，联合厚度约 3.41 µm，相对多光束联合结果偏低约 5%。"
        "色散坐标多峰谷模型对硅给出约3.400 µm的候选值，但峰谷差为4.771%、留段CV为3.168%，"
        "且与Airy结果相差5.419%，故按门控回退多光束模型。"
        "碳化硅因不满足四项多光束证据门控，仍采用问题二的双光束机理与色散坐标名义厚度。",
    )

    add(doc, "6.5 总结", kind="h2")
    add(
        doc,
        "问题三建立了多光束可观测条件、Airy 模型与四项证据门控。"
        "结论：附件三、四存在可观察多光束干涉，色散坐标极值候选未通过稳定性门控，故d_Si≈3.59 µm；"
        "附件一、二多光束证据不足，SiC不作多光束修正，采用问题二色散坐标结果d_SiC≈7.45 µm。",
    )

    return doc


def main():
    doc = build()
    saved = OUT
    try:
        doc.save(str(OUT))
    except PermissionError:
        saved = OUT_ALT
        doc.save(str(OUT_ALT))
    MD_OUT.write_text(
        """# 详细版（按既定标题重写）

## 一二级标题（未改名）
- 问题重述 / 1.1 问题背景 / 1.2 问题提出
- 问题分析 / 2.2问题一分析 / 2.2问题二分析 / 2.2问题三分析
- 符号说明
- 问题一 / 4.1数据预处理 / 4.2模型建立 / 4,3 模型求解 / 4.4 结果分析 / 4.5 总结
- 问题二 / 5.1–5.5
- 问题三 / 6.1–6.5

## 新增三级标题（仅用于细分步骤）
- 4.2.1–4.2.4（折射、光程相位、反射率、厚度公式）
- 5.2.1–5.2.2（条纹测厚、可靠性）
- 6.2.1–6.2.3（可观测条件、Airy、证据门控）

## 公式写法
- 步骤文字详细；公式间用「由……得到」简要承接，不展开代数推导细节。

文件：`2025B论文_问题一二三_详细版.docx`（占用时另存 `*_公式修订.docx`）
""",
        encoding="utf-8",
    )
    print("OUT", saved.as_posix())


if __name__ == "__main__":
    main()
