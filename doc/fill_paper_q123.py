# -*- coding: utf-8 -*-
"""在原 2025B论文.docx 上填充问题一二三，尽量保留已有图片。"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_DIR = Path(__file__).resolve().parent
SRC = DOC_DIR / "2025B论文.docx"
BACKUP = DOC_DIR / "2025B论文_填充前备份.docx"
OUT_COPY = DOC_DIR / "2025B论文_问题一二三.docx"
MD_OUT = DOC_DIR / "2025B论文_问题一二三.md"


def set_run_font(run, *, east_asia: str, ascii_font: str = "Times New Roman", size_pt: float, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def style_paragraph(p, *, level: str = "body", first_indent: bool = True, center: bool = False):
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    if level == "h1":
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(12)
    elif level in {"h2", "h3", "caption", "formula", "title"}:
        pf.first_line_indent = Pt(0)
        if level == "h2":
            pf.space_before = Pt(8)
        if level in {"caption", "formula", "title"}:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.first_line_indent = Cm(0.74) if first_indent else Pt(0)


def add_para(doc, text: str, *, level: str = "body", first_indent: bool = True, center: bool = False):
    p = doc.add_paragraph()
    style_paragraph(p, level=level, first_indent=first_indent, center=center)
    if level == "title":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=15, bold=True)
    elif level == "h1":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=15, bold=True)
    elif level == "h2":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=12, bold=True)
    elif level == "h3":
        set_run_font(p.add_run(text), east_asia="黑体", size_pt=10.5, bold=True)
    elif level == "caption":
        set_run_font(p.add_run(text), east_asia="宋体", size_pt=10.5, bold=True)
    else:
        set_run_font(p.add_run(text), east_asia="宋体", size_pt=10.5)
    return p


def add_table(doc, headers, rows, caption: str):
    add_para(doc, caption, level="caption")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), east_asia="黑体", size_pt=10.5, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(val), east_asia="宋体", size_pt=10.5)
    doc.add_paragraph()


def clear_after_analysis(doc: Document):
    """保留标题、问题重述、问题分析；从“符号说明/问题一”起清空。"""
    cut_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t in {"符号说明", "问题一"} or t.startswith("4.1"):
            cut_idx = i
            break
    if cut_idx is None:
        raise RuntimeError("未找到符号说明/问题一起点")
    # 修正分析小节编号
    for p in doc.paragraphs:
        if p.text.strip() == "2.2问题一分析":
            p.text = ""
            set_run_font(p.add_run("2.1 问题一分析"), east_asia="黑体", size_pt=12, bold=True)
            style_paragraph(p, level="h2")
        elif p.text.strip() == "2.2问题二分析":
            p.text = ""
            set_run_font(p.add_run("2.2 问题二分析"), east_asia="黑体", size_pt=12, bold=True)
            style_paragraph(p, level="h2")
        elif p.text.strip() == "2.2问题三分析":
            p.text = ""
            set_run_font(p.add_run("2.3 问题三分析"), east_asia="黑体", size_pt=12, bold=True)
            style_paragraph(p, level="h2")
        elif p.text.strip() == "问题分析":
            p.text = ""
            set_run_font(p.add_run("二、问题分析"), east_asia="黑体", size_pt=15, bold=True)
            style_paragraph(p, level="h1")
        elif p.text.strip() == "问题重述":
            p.text = ""
            set_run_font(p.add_run("一、问题重述"), east_asia="黑体", size_pt=15, bold=True)
            style_paragraph(p, level="h1")

    body = doc.element.body
    # 删除 cut_idx 及之后的段落（保留 sectPr）
    paras = list(doc.paragraphs)
    for p in paras[cut_idx:]:
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    return cut_idx


def append_body(doc: Document):
    add_para(doc, "三、模型假设", level="h1")
    for i, item in enumerate(
        [
            "外延层上下表面局部平行，分析光斑内厚度可视为常数。",
            "空气折射率取 n0=1，入射角相对样品表面法线定义。",
            "主厚度反演在弱吸收透明波段进行；碳化硅约 797–1000 cm⁻¹ 的强声子吸收带不参与主估计。",
            "双/多光束条纹反演阶段，SiC 与 Si 的常折射率基线分别取 n_SiC=2.55、n_Si=3.42；该假设下的厚度为条件估计。",
            "光谱基线与条纹振幅可随波数缓慢变化，但同一材料在两个入射角下共享同一物理厚度。",
            "多光束判定需谐波比、有效反射率与模型改善三类证据同时通过，避免仅凭残差下降误报。",
        ],
        1,
    ):
        add_para(doc, f"{i}. {item}")

    add_para(doc, "四、符号说明", level="h1")
    add_para(doc, "正文主要符号见表1。未列入的中间变量在首次出现处给出定义。")
    add_table(
        doc,
        ["符号", "含义", "单位"],
        [
            ["θ₀", "空气中入射角", "°"],
            ["θ₁", "外延层内折射角", "°"],
            ["n / ñ", "折射率 / 复折射率", "—"],
            ["ν̃", "波数", "cm⁻¹"],
            ["d", "外延层厚度", "µm"],
            ["δ", "一次往返相位差", "rad"],
            ["Δν̃", "相邻同类极值波数间隔", "cm⁻¹"],
            ["R", "反射率", "— 或 %"],
            ["Rᵢ", "有效界面反射率", "—"],
            ["A₂/A₁", "二次谐波与基频幅值比", "—"],
            ["ΔAICc", "双光束相对多光束的 AICc 差值", "—"],
        ],
        "表1  主要符号说明",
    )

    # 问题一
    add_para(doc, "五、问题一：双光束干涉厚度模型", level="h1")
    add_para(doc, "5.1 数据特征与预处理框架", level="h2")
    add_para(
        doc,
        "四份附件各含约 7469 个采样点，波数范围约 399.7–4000.1 cm⁻¹，中位间隔约 0.482 cm⁻¹。"
        "程序对原始表做有限值过滤、波数升序与去重，并剔除首点可疑零值，得到可分析序列。"
        "问题一只建立解析模型，不输出具体厚度；但预处理与光学几何约定将贯穿问题二、三。",
    )
    add_para(doc, "5.2 模型建立", level="h2")
    add_para(doc, "5.2.1 折射角与光程差", level="h3")
    add_para(doc, "设空气入射角为 θ₀，外延层折射率为 n，折射角为 θ₁。由 Snell 定律")
    add_para(doc, "sin θ₀ = n sin θ₁，  cos θ₁ = √(1 − sin²θ₀ / n²)。", level="formula")
    add_para(
        doc,
        "其中 cos θ₁ 取物理可行的正分支。膜内一次往返光程差为 ΔL = 2 n d cos θ₁。"
        "以波数 ν̃=1/λ 表示时，两束反射光的相位差写为",
        first_indent=False,
    )
    add_para(doc, "δ(ν̃) = 4π d ν̃ n cos θ₁ + φᵣ。", level="formula")
    add_para(
        doc,
        "其中 φᵣ 为两界面反射引入的近似固定相位。它改变峰谷对应的干涉级次归属，"
        "但不改变相邻同类极值之间的波数间隔，因此厚度反演可主要依赖间隔信息。",
        first_indent=False,
    )
    add_para(doc, "5.2.2 双光束反射率与厚度公式", level="h3")
    add_para(
        doc,
        "在每个界面仅发生一次反射/透射的双光束假设下，总反射场可视为两束光叠加。"
        "若慢变振幅包络记为 A(ν̃)、B(ν̃)，则反射率可写为",
    )
    add_para(doc, "R(ν̃) = A(ν̃) + B(ν̃) cos δ(ν̃)。", level="formula")
    add_para(
        doc,
        "定义光学相位坐标 g(ν̃,θ₀)=ν̃ √(n²−sin²θ₀)。"
        "若两个同类极值跨越 Δm 个干涉级次，则",
        first_indent=False,
    )
    add_para(doc, "d = Δm / {2 [g(ν̃_b,θ₀) − g(ν̃_a,θ₀)]}。", level="formula")
    add_para(
        doc,
        "当窄波段内折射率近似为常数、且取相邻同类极值（Δm=1）时，得到问题一的核心解析关系：",
        first_indent=False,
    )
    add_para(doc, "d = 1 / [2 Δν̃ √(n² − sin²θ₀)]。", level="formula")
    add_para(
        doc,
        "其中 d 以 µm、Δν̃ 以 cm⁻¹ 计时，右侧再乘 10⁴。"
        "当 θ₀→0 时，公式退化为 d=1/(2nΔν̃)，可作为推导正确性的退化检验。",
        first_indent=False,
    )
    add_para(doc, "5.3 模型求解思路", level="h2")
    add_para(
        doc,
        "问题一本身是机理建模。后续数值实现时，厚度求解按“粗估—稳健极值—精修”三步展开："
        "（1）短窗 Savitzky–Golay 去噪、长窗提取慢变基线，得到条纹残差；"
        "（2）在波数域做 FFT，于 1–20 µm 物理范围内获得厚度粗值；"
        "（3）按粗周期设置峰距，提取峰/谷，并以 Theil–Sen 回归估计间距，抑制漏峰与误峰；"
        "必要时在稳健解附近做有界相位精修。该流程在问题二中具体落地。",
    )
    add_para(doc, "5.4 结果分析", level="h2")
    add_para(
        doc,
        "问题一给出的是可检验的解析映射而非数值厚度。其直接用途有三："
        "第一，把斜入射几何正确写入相位，避免用正入射公式硬套十度、十五度数据；"
        "第二，说明厚度主要由同类极值间隔决定，为峰谷稳健估计提供理论依据；"
        "第三，明确常折射率假设的适用范围，为问题二选择透明波段、问题三判断多光束修正提供边界。",
    )
    add_para(doc, "5.5 小结", level="h2")
    add_para(
        doc,
        "本节建立了斜入射双光束干涉厚度模型，得到常折射率下厚度与相邻同类极值波数间隔的解析公式，"
        "并给出与后续数值实现衔接的求解框架。该模型构成问题二碳化硅测厚与问题三模型分叉的共同起点。",
    )

    # 问题二
    add_para(doc, "六、问题二：碳化硅外延层厚度反演与可靠性评价", level="h1")
    add_para(doc, "6.1 数据预处理与波段选择", level="h2")
    add_para(
        doc,
        "附件1、2分别对应同一块碳化硅晶圆在 10° 与 15° 入射角下的反射谱。"
        "在问题一公式基础上，首先避开强声子吸收带，对候选透明窗评分并自动选带；"
        "最终用于条纹拟合的主波段约为 1200–4000 cm⁻¹。"
        "短窗滤波抑制点噪声，长窗估计慢变基线，残差谱供 FFT、峰谷识别与全谱精修使用。"
        "需指出：附件2最大反射率约 102.74%，说明绝对反射率可能存在仪器标定问题；"
        "双光束厚度主要由峰位/间距决定，故该异常主要记入审计，不强制裁剪到 100%。",
    )
    add_para(doc, "6.2 模型建立", level="h2")
    add_para(doc, "6.2.1 双光束稳健测厚模型", level="h3")
    add_para(
        doc,
        "对选定波段内的残差条纹，采用问题一公式。设峰序列与谷序列分别得到间距估计，"
        "再换算为厚度 d_peak、d_valley；最终双光束采用值取二者稳健综合后的精修结果 d_two。"
        "折射率取 n_SiC=2.55。该模型自由度低、物理解释清楚，适合作为碳化硅主结果。",
    )
    add_para(doc, "6.2.2 不确定度与双角度一致性", level="h3")
    add_para(
        doc,
        "统计不确定度由峰/谷间距重采样得到条件 95% 区间；系统不确定度则来自折射率假设。"
        "由 d ∝ (n²−sin²θ₀)^(−1/2) 可知，近法向时 Δd/d≈−Δn/n，"
        "例如 n 偏差 1% 将引起约 1% 的厚度反向偏差。"
        "同一晶圆应满足两角度共享厚度，故定义相对差",
    )
    add_para(doc, "ε = |d₁₀ − d₁₅| / [(d₁₀+d₁₅)/2] × 100%。", level="formula")
    add_para(doc, "并以重采样标准差的逆方差加权得到联合厚度，作为问题二最终报告值。", first_indent=False)
    add_para(doc, "6.3 模型求解", level="h2")
    add_para(
        doc,
        "对附件1、2分别执行：选带 → 预处理 → FFT 粗估 → 峰谷 Theil–Sen → 有界精修 → 间距 bootstrap。"
        "随后计算双角度相对差与加权联合厚度。程序入口为完整流水线，保证两角度使用同一套阈值与随机种子规则。",
    )
    add_para(doc, "6.4 结果分析", level="h2")
    add_table(
        doc,
        ["入射角", "峰序列厚度/µm", "谷序列厚度/µm", "双光束采用值/µm", "条件95%区间/µm"],
        [
            ["10°", "7.839", "7.926", "7.883", "[7.677, 8.152]"],
            ["15°", "7.811", "7.772", "7.791", "[7.560, 8.051]"],
        ],
        "表2  碳化硅双光束厚度反演结果",
    )
    add_para(
        doc,
        "由表2可见，峰、谷两套估计接近，说明极值识别稳定。"
        "两角度采用值的相对差约为 1.16%，小于 2% 的工程一致性门槛；"
        "按重采样标准差加权，得到",
        first_indent=False,
    )
    add_para(doc, "d_SiC ≈ 7.83 µm。", level="formula")
    add_para(
        doc,
        "可靠性评价可从四方面概括："
        "（1）内部一致性：峰谷接近、双角度相对差约 1%；"
        "（2）采样能力：条纹间距约 240–250 cm⁻¹，而采样间隔约 0.482 cm⁻¹，每条条纹约有数百点；"
        "（3）统计误差：表中区间反映噪声与局部条纹不均匀；"
        "（4）系统误差：折射率来源误差不能由重采样消除，需在正式应用中用晶圆对应光学常数更新。"
        "此外，问题三将证明现有证据不足以支持对 SiC 做多光束修正，故问题二主结果保留双光束模型。",
        first_indent=False,
    )
    add_para(doc, "6.5 小结", level="h2")
    add_para(
        doc,
        "针对附件1、2，本文在自动透明波段内建立双光束稳健测厚流程，得到碳化硅外延层厚度约 7.83 µm，"
        "双角度相对差约 1.16%，并给出条件置信区间与折射率系统误差说明。该结果作为后续多光束诊断的对照基线。",
    )

    # 问题三
    add_para(doc, "七、问题三：多光束干涉条件、硅片测厚与碳化硅修正判定", level="h1")
    add_para(doc, "7.1 数据预处理", level="h2")
    add_para(
        doc,
        "附件3、4为硅片在 10°、15° 的反射谱；附件1、2继续用于检验碳化硅是否需要多光束修正。"
        "预处理与问题二一致：选带、双尺度滤波与去基线。多光束诊断在残差条纹上进行，"
        "同时用 Airy 模型对完整峰形做有界拟合，以便与双光束模型比较信息准则与误差。",
    )
    add_para(doc, "7.2 模型建立", level="h2")
    add_para(doc, "7.2.1 多光束可观测必要条件", level="h3")
    add_para(
        doc,
        "外延层内第 j 次往返相对前一束多乘复因子 q=r₁₀ r₁₂ exp(−2αd/cosθ₁) exp(iδ)。"
        "仅“存在后续反射”并不足够，还需同时满足：界面反射足够强、往返吸收不过强、"
        "光源相干长度覆盖相关光程差、探测面空间重合与偏振可干涉、仪器分辨率小于自由光谱范围，"
        "以及膜厚均匀导致的相位展宽不足以抹平条纹。据此，可观测性必须由数据证据检验，而非先验假定。",
    )
    add_para(doc, "7.2.2 Airy/Fresnel 多光束反射模型", level="h3")
    add_para(doc, "对 s、p 偏振，单层多次反射总振幅可写为")
    add_para(
        doc,
        "r_tot = (r₀₁ + r₁₂ e^{2iβ}) / (1 + r₀₁ r₁₂ e^{2iβ})，  β = 2π d ν̃ n cos θ₁。",
        level="formula",
    )
    add_para(
        doc,
        "非偏振反射率取二者强度平均。对称无吸收情形化为 Airy 形式，精细度由有效界面反射率 Rᵢ 控制。"
        "数值实现中，非线性参数为 (d, Rᵢ, 相位偏置)，线性增益与慢变基线用最小二乘剖面消元；"
        "可选用差分进化全局搜索再局部精修。",
        first_indent=False,
    )
    add_para(doc, "7.2.3 四项证据联合判定", level="h3")
    add_para(
        doc,
        "为避免“多参数必然降残差”造成的误判，规定须同时满足："
        "谐波比 A₂/A₁≥0.08、有效反射率 Rᵢ≥0.12、相对双光束的 RMSE 改善≥2%，以及 ΔAICc≥10。"
        "其中有效反射率可被未建模基线部分吸收，故不能单独作为证据；谐波与信息准则是关键补充。",
    )
    add_para(doc, "7.3 模型求解", level="h2")
    add_para(
        doc,
        "对四份光谱分别拟合双光束与 Airy 多光束模型，计算谐波频谱与四项指标。"
        "若判定为可观测多光束，则采用多光束厚度；否则回退双光束厚度。"
        "最后对硅片两角度结果做一致性检验与逆方差加权融合。",
    )
    add_para(doc, "7.4 结果分析", level="h2")
    add_table(
        doc,
        ["数据集", "A₂/A₁", "有效反射率", "RMSE改善/%", "ΔAICc", "判定"],
        [
            ["SiC 10°", "0.019", "0.002", "23.2", "3072", "证据不足"],
            ["SiC 15°", "0.037", "0.005", "25.9", "3480", "证据不足"],
            ["Si 10°", "0.196", "0.410", "11.9", "1475", "可观测多光束"],
            ["Si 15°", "0.222", "0.422", "8.65", "1051", "可观测多光束"],
        ],
        "表3  多光束四项证据诊断结果",
    )
    add_para(
        doc,
        "表3表明：硅片两角度谐波比约 0.20–0.22，有效反射率约 0.41–0.42，"
        "模型改善与 ΔAICc 同时显著，故判定附件3、4存在可观察的多光束干涉。"
        "碳化硅虽然 Airy 拟合可使 RMSE 下降，但谐波比仅约 0.02–0.04，有效反射率接近下界，"
        "不满足联合证据门槛，因此主结果不作多光束修正。",
        first_indent=False,
    )
    add_table(
        doc,
        ["入射角", "双光束厚度/µm", "Airy多光束厚度/µm", "最终采用模型", "采用厚度/µm"],
        [
            ["Si 10°", "3.403", "3.598", "多光束", "3.598"],
            ["Si 15°", "3.410", "3.584", "多光束", "3.584"],
            ["SiC 10°", "7.883", "8.449", "双光束", "7.883"],
            ["SiC 15°", "7.791", "8.333", "双光束", "7.791"],
        ],
        "表4  双光束与多光束厚度对比及最终采用值",
    )
    add_para(doc, "硅片两角度多光束厚度相对差约 0.38%，逆方差加权联合结果为", first_indent=False)
    add_para(doc, "d_Si ≈ 3.59 µm。", level="formula")
    add_para(
        doc,
        "若硅片仍用双光束模型，联合厚度约 3.41 µm，相对多光束结果偏低约 5%，"
        "说明忽略多次反射会引入不可忽略的模型误差。"
        "对碳化硅，Airy 厚度虽系统性偏高，但因谐波证据不足，仍以双光束 7.83 µm 为最终报告值。"
        "理想常折射率无吸收情形下，仅由准确极值间隔计算的厚度理论位置可与 Airy 峰位相容；"
        "实际偏差主要来自峰形非余弦、采样与基线耦合，因此“能否观测多光束”与“厚度点估计是否必须修正”应分开讨论。",
        first_indent=False,
    )
    add_para(doc, "7.5 小结", level="h2")
    add_para(
        doc,
        "本节给出多光束可观测条件与 Airy 拟合—四项证据门控流程。"
        "结论为：附件3、4硅片存在显著多光束干涉，外延层厚度约 3.59 µm；"
        "附件1、2碳化硅现有证据不足，不作多光束修正，厚度仍取问题二的双光束结果约 7.83 µm。",
    )

    add_para(doc, "八、问题一二三结果汇总", level="h1")
    add_table(
        doc,
        ["问题", "对象", "采用模型", "关键结果"],
        [
            ["一", "机理模型", "斜入射双光束干涉", "给出 d–Δν̃ 解析关系"],
            ["二", "SiC 附件1/2", "双光束稳健极值", "d≈7.83 µm，相对差≈1.16%"],
            ["三", "Si 附件3/4", "Airy 多光束", "d≈3.59 µm，相对差≈0.38%"],
            ["三", "SiC 修正判定", "证据门控回退双光束", "不作多光束修正"],
        ],
        "表5  问题一二三主要结论汇总",
    )
    add_para(
        doc,
        "上述数值均基于题面常折射率基线与当前附件条件。若补充晶型、掺杂与温度对应的权威光学常数，"
        "可在同一程序框架下更新有效数字；色散与载流子浓度耦合分析可作为可靠性扩展，但不改变本节主结论的报告口径。",
        first_indent=False,
    )


def write_markdown():
    MD_OUT.write_text(
        """# 2025B 论文：问题一二三正文（与 Word 同步）

## 关键结果（与 output/thickness_summary.csv 一致）

- 问题一：斜入射双光束解析式 `d = 1 / [2 Δν̃ √(n²-sin²θ0)]`
- 问题二：SiC 双光束，10°=7.883 µm，15°=7.791 µm，相对差≈1.16%，加权 **d_SiC≈7.83 µm**
- 问题三：Si 可观测多光束，10°=3.598 µm，15°=3.584 µm，相对差≈0.38%，加权 **d_Si≈3.59 µm**；SiC 证据不足，不作多光束修正

详细正文见 `2025B论文_问题一二三.docx` / 原 `2025B论文.docx`。
""",
        encoding="utf-8",
    )


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    # 原文件可能被 Word 占用：可读则备份；写回失败时仅输出副本
    try:
        source_bytes = SRC.read_bytes()
        if not BACKUP.exists():
            BACKUP.write_bytes(source_bytes)
    except PermissionError:
        if BACKUP.exists():
            source_bytes = BACKUP.read_bytes()
        else:
            raise RuntimeError("原论文被占用且无备份，请先关闭 Word 后重试")

    # 优先从备份构建，避免半成品覆盖
    build_from = BACKUP if BACKUP.exists() else SRC
    doc = Document(str(build_from))
    clear_after_analysis(doc)
    append_body(doc)

    doc.save(str(OUT_COPY))
    write_markdown()

    try:
        doc.save(str(SRC))
        status = "updated_original"
    except PermissionError:
        status = "original_locked_use_copy"

    print("STATUS", status)
    print("COPY", OUT_COPY.as_posix())
    print("TABLES", len(doc.tables))
    print("PARAS", len(doc.paragraphs))


if __name__ == "__main__":
    main()
